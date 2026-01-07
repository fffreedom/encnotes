#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
导出管理器 - 支持导出为PDF、Word、Markdown格式
"""

import os
import re
from pathlib import Path
from PyQt6.QtGui import QTextDocument, QPageSize, QPageLayout
from PyQt6.QtCore import QMarginsF, QSizeF
from PyQt6.QtPrintSupport import QPrinter
from datetime import datetime


class ExportManager:
    """导出管理器类"""
    
    def __init__(self):
        self.export_dir = Path.home() / "Documents" / "encnotes导出"
        self.export_dir.mkdir(parents=True, exist_ok=True)
        
    def export_to_pdf(self, note_title, note_content_html):
        """
        导出为PDF格式
        
        Args:
            note_title: 笔记标题
            note_content_html: 笔记HTML内容
            
        Returns:
            导出的文件路径，失败返回None
        """
        try:
            # 生成文件名
            safe_title = self._sanitize_filename(note_title)
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{safe_title}_{timestamp}.pdf"
            filepath = self.export_dir / filename
            
            # 创建打印机对象
            printer = QPrinter(QPrinter.PrinterMode.HighResolution)
            printer.setOutputFormat(QPrinter.OutputFormat.PdfFormat)
            printer.setOutputFileName(str(filepath))
            
            # 设置页面大小和边距
            page_layout = QPageLayout()
            page_layout.setPageSize(QPageSize(QPageSize.PageSizeId.A4))
            page_layout.setOrientation(QPageLayout.Orientation.Portrait)
            page_layout.setMargins(QMarginsF(20, 20, 20, 20))
            printer.setPageLayout(page_layout)
            
            # 创建文档并打印
            document = QTextDocument()
            document.setHtml(note_content_html)
            document.print_(printer)
            
            return str(filepath)
            
        except Exception as e:
            print(f"导出PDF失败: {e}")
            return None
            
    def export_to_word(self, note_title, note_content_html):
        """
        导出为Word格式（.docx）
        
        Args:
            note_title: 笔记标题
            note_content_html: 笔记HTML内容
            
        Returns:
            导出的文件路径，失败返回None
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from bs4 import BeautifulSoup
            
            # 生成文件名
            safe_title = self._sanitize_filename(note_title)
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{safe_title}_{timestamp}.docx"
            filepath = self.export_dir / filename
            
            # 创建Word文档
            doc = Document()
            
            # 添加标题
            title = doc.add_heading(note_title, level=1)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加创建时间
            time_para = doc.add_paragraph()
            time_para.add_run(f"导出时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            time_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            doc.add_paragraph()  # 空行
            
            # 解析HTML内容
            soup = BeautifulSoup(note_content_html, 'html.parser')
            
            # 提取文本内容
            text_content = soup.get_text()
            
            # 按段落分割
            paragraphs = text_content.split('\n')
            
            for para_text in paragraphs:
                if para_text.strip():
                    para = doc.add_paragraph(para_text.strip())
                    
            # 保存文档
            doc.save(str(filepath))
            
            return str(filepath)
            
        except ImportError:
            print("导出Word需要安装 python-docx 和 beautifulsoup4 库")
            print("请运行: pip install python-docx beautifulsoup4")
            return None
        except Exception as e:
            print(f"导出Word失败: {e}")
            return None
            
    def export_to_markdown(self, note_title, note_content_html):
        """
        导出为Markdown格式
        
        Args:
            note_title: 笔记标题
            note_content_html: 笔记HTML内容
            
        Returns:
            导出的文件路径，失败返回None
        """
        try:
            from bs4 import BeautifulSoup
            import html2text
            
            # 生成文件名
            safe_title = self._sanitize_filename(note_title)
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{safe_title}_{timestamp}.md"
            filepath = self.export_dir / filename
            
            # 转换HTML为Markdown
            h = html2text.HTML2Text()
            h.ignore_links = False
            h.ignore_images = False
            h.ignore_emphasis = False
            h.body_width = 0  # 不自动换行
            
            markdown_content = h.handle(note_content_html)
            
            # 添加标题和元数据
            full_content = f"# {note_title}\n\n"
            full_content += f"*导出时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}*\n\n"
            full_content += "---\n\n"
            full_content += markdown_content
            
            # 保存文件
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(full_content)
                
            return str(filepath)
            
        except ImportError:
            print("导出Markdown需要安装 html2text 和 beautifulsoup4 库")
            print("请运行: pip install html2text beautifulsoup4")
            return None
        except Exception as e:
            print(f"导出Markdown失败: {e}")
            return None
            
    def export_to_html(self, note_title, note_content_html):
        """
        导出为HTML格式
        
        Args:
            note_title: 笔记标题
            note_content_html: 笔记HTML内容
            
        Returns:
            导出的文件路径，失败返回None
        """
        try:
            # 生成文件名
            safe_title = self._sanitize_filename(note_title)
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{safe_title}_{timestamp}.html"
            filepath = self.export_dir / filename
            
            # 创建完整的HTML文档
            html_template = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{note_title}</title>
    <style>
        body {{
            font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, "Helvetica Neue", Arial, sans-serif;
            max-width: 800px;
            margin: 0 auto;
            padding: 40px 20px;
            line-height: 1.6;
            color: #333;
        }}
        h1 {{
            text-align: center;
            color: #2c3e50;
            border-bottom: 2px solid #3498db;
            padding-bottom: 10px;
        }}
        .meta {{
            text-align: right;
            color: #7f8c8d;
            font-size: 0.9em;
            margin-bottom: 30px;
        }}
        .content {{
            margin-top: 30px;
        }}
    </style>
</head>
<body>
    <h1>{note_title}</h1>
    <div class="meta">导出时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</div>
    <hr>
    <div class="content">
        {note_content_html}
    </div>
</body>
</html>"""
            
            # 保存文件
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(html_template)
                
            return str(filepath)
            
        except Exception as e:
            print(f"导出HTML失败: {e}")
            return None
            
    def _sanitize_filename(self, filename):
        """
        清理文件名，移除非法字符
        
        Args:
            filename: 原始文件名
            
        Returns:
            清理后的文件名
        """
        # 移除非法字符
        filename = re.sub(r'[<>:"/\\|?*]', '', filename)
        # 限制长度
        filename = filename[:50]
        # 如果为空，使用默认名称
        if not filename.strip():
            filename = "未命名笔记"
        return filename
        
    def get_export_directory(self):
        """获取导出目录路径"""
        return str(self.export_dir)
